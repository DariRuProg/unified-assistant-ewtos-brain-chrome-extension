# @author Dario | ewtos.com
import io


def extract_text(data: bytes, filename: str = "") -> str:
    """Extrahiert Text aus PDF- oder Textdaten. Gibt Markdown-String zurück."""
    fname = (filename or "").lower()
    if fname.endswith(".pdf"):
        info = _extract_pdf(data)
        return info["text"]
    return data.decode("utf-8", errors="replace")


def extract_info(data: bytes, filename: str = "") -> dict:
    """Wie extract_text, liefert zusätzlich Metadaten (title, author, pages).

    Returns:
        {"text": str, "title": str, "author": str, "pages": int}
    """
    fname = (filename or "").lower()
    if fname.endswith(".pdf"):
        return _extract_pdf(data)
    if fname.endswith(".docx"):
        return _extract_docx(data)
    return {"text": data.decode("utf-8", errors="replace"), "title": "", "author": "", "pages": 0}


def _extract_docx(data: bytes) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx nicht installiert. `pip install python-docx` ausführen.")
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"DOCX konnte nicht gelesen werden: {e}")
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    props = doc.core_properties
    return {
        "text": "\n\n".join(parts),
        "title": (props.title or "").strip(),
        "author": (props.author or "").strip(),
        "pages": 0,
    }


def _extract_pdf(data: bytes) -> dict:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber nicht installiert. `pip install pdfplumber` ausführen.")
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            meta = pdf.metadata or {}
            pages_text = []
            for page in pdf.pages:
                try:
                    text = page.extract_text()
                except Exception:
                    text = None
                if text and text.strip():
                    pages_text.append(text.strip())
            text = "\n\n---\n\n".join(pages_text)
    except Exception as e:
        raise ValueError(f"PDF konnte nicht gelesen werden: {e}")
    return {
        "text": text,
        "title": (meta.get("Title") or "").strip(),
        "author": (meta.get("Author") or "").strip(),
        "pages": len(pages_text),
    }
